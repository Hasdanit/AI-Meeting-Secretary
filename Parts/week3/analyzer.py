import google.generativeai as genai
import os
from dotenv import load_dotenv
from docx import Document
from datetime import datetime
import re
from notion_integration import add_task_to_notion

load_dotenv()

api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise ValueError("–û—à–∏–±–∫–∞: GEMINI_API_KEY –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω.")

genai.configure(api_key=api_key)
model = genai.GenerativeModel("gemini-1.5-flash")

def analyze_text(text: str):
    prompt = (
        "–¢—ã –ò–ò-—Å–µ–∫—Ä–µ—Ç–∞—Ä—å. –ü—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä—É–π —Ç–µ–∫—Å—Ç –≤—Å—Ç—Ä–µ—á–∏, –∫–æ—Ç–æ—Ä—ã–π –≤–∫–ª—é—á–∞–µ—Ç –≤—Ä–µ–º–µ–Ω–Ω—ã–µ –º–µ—Ç–∫–∏ –¥–ª—è –∫–∞–∂–¥–æ–≥–æ —Å–µ–≥–º–µ–Ω—Ç–∞. "
        "–°–≥–µ–Ω–µ—Ä–∏—Ä—É–π –∫—Ä–∞—Ç–∫–æ–µ —Ä–µ–∑—é–º–µ –≤—Å—Ç—Ä–µ—á–∏, —Ö—Ä–æ–Ω–æ–ª–æ–≥–∏—é –æ–±—Å—É–∂–¥–∞–µ–º—ã—Ö —Ç–µ–º —Å —É–∫–∞–∑–∞–Ω–∏–µ–º –≤—Ä–µ–º–µ–Ω–Ω—ã—Ö –º–µ—Ç–æ–∫ –∏ —Å–ø–∏—Å–æ–∫ –∑–∞–¥–∞—á. "
        "–ò—Å–ø–æ–ª—å–∑—É–π –ø—Ä–µ–¥–æ—Å—Ç–∞–≤–ª–µ–Ω–Ω—ã–µ –≤—Ä–µ–º–µ–Ω–Ω—ã–µ –º–µ—Ç–∫–∏ –¥–ª—è –æ–±–æ–∑–Ω–∞—á–µ–Ω–∏—è –≤—Ä–µ–º–µ–Ω–∏ –Ω–∞—á–∞–ª–∞ –∫–∞–∂–¥–æ–π —Ç–µ–º—ã –∏–ª–∏ —Å–æ–±—ã—Ç–∏—è.\n\n"
        "–§–æ—Ä–º–∞—Ç –æ—Ç–≤–µ—Ç–∞:\n"
        "–†–µ–∑—é–º–µ:\n[—Ç–µ–∫—Å—Ç —Ä–µ–∑—é–º–µ]\n\n"
        "–¢–∞–π–º–ª–∞–π–Ω:\n- [–≤—Ä–µ–º—è] - [—Ç–µ–º–∞]\n\n"
        "–ó–∞–¥–∞—á–∏:\n- [–∑–∞–¥–∞—á–∞]\n\n"
        "–¢–µ–∫—Å—Ç –≤—Å—Ç—Ä–µ—á–∏:\n" + text
    )

    try:
        response = model.generate_content(prompt)
        analysis = response.text

        summary = re.search(r'–†–µ–∑—é–º–µ:\s*(.*?)\s*–¢–∞–π–º–ª–∞–π–Ω:', analysis, re.DOTALL).group(1).strip()
        timeline = re.search(r'–¢–∞–π–º–ª–∞–π–Ω:\s*(.*?)\s*–ó–∞–¥–∞—á–∏:', analysis, re.DOTALL).group(1).strip()
        tasks = re.search(r'–ó–∞–¥–∞—á–∏:\s*(.*)', analysis, re.DOTALL).group(1).strip().split('\n')

        return summary, timeline, tasks
    except Exception as e:
        print(f"–ü—Ä–æ–∏–∑–æ—à–ª–∞ –æ—à–∏–±–∫–∞ –ø—Ä–∏ –∞–Ω–∞–ª–∏–∑–µ —Ç–µ–∫—Å—Ç–∞: {e}")
        return None, None, None

def save_to_docx(filename_base: str, summary: str, timeline: str, tasks: list, raw_text: str):
    doc = Document()
    doc.add_heading("–û—Ç—á—ë—Ç –ø–æ –≤—Å—Ç—Ä–µ—á–µ", level=0)
    doc.add_paragraph(f"–î–∞—Ç–∞ –∏ –≤—Ä–µ–º—è: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_heading("üîπ –†–µ–∑—é–º–µ:", level=1)
    doc.add_paragraph(summary)
    doc.add_heading("üîπ –¢–∞–π–º–ª–∞–π–Ω:", level=1)
    for line in timeline.split('\n'):
        doc.add_paragraph(line)
    doc.add_heading("üîπ –ó–∞–¥–∞—á–∏:", level=1)
    for task in tasks:
        doc.add_paragraph(task, style='List Bullet')
    doc.add_heading("üîπ –ü–æ–ª–Ω–∞—è —Å—Ç–µ–Ω–æ–≥—Ä–∞–º–º–∞:", level=1)
    doc.add_paragraph(raw_text)

    output_path = os.path.join("recordings", f"{filename_base}.docx")
    doc.save(output_path)
    print(f"[‚úî] Word-—Ñ–∞–π–ª —Å–æ—Ö—Ä–∞–Ω—ë–Ω: {output_path}")

    for task in tasks:
        add_task_to_notion(task.strip('- ').strip())